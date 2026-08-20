import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
from collections import Counter
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from triage_assistant import render_triage_ui
from debugger import render_debugger_ui
from debug_analytics import render_analytics_ui
from data_explorer import render_data_explorer
from schematic_viewer import render_schematic_viewer
from debug_workspace import render_debug_workspace
import convert_jupiter_csv as convert_jupiter
from program_config import (
    get_program_list, register_program, get_selected_program,
    set_selected_program, load_registry, get_program_dir
)
import theme
from theme import (
    PURPLE, PURPLE_BRIGHT, PURPLE_DEEP, GREEN, MAGENTA, WARNING,
    PANEL, BG_ELEV, BORDER, TEXT, TEXT_MUTED, TEXT_FAINT,
    CHART_SEQUENCE, CONTINUOUS_SCALE, APP_NAME, APP_VERSION,
)

# Page configuration
st.set_page_config(
    page_title=f"{APP_NAME}",
    page_icon=theme.page_icon(),
    layout="wide",
    initial_sidebar_state="expanded"
)

# Global theme (DM Sans + Kiro-purple palette). Runs on every rerun.
theme.inject_theme()

@st.cache_data(show_spinner=False)
def _load_dashboard_dataframe(file_bytes):
    """Parse and clean the uploaded CSV. Cached on file content so it only
    re-parses when a different file is uploaded."""
    df = pd.read_csv(io.BytesIO(file_bytes))
    # Filter out "Won't do" cases from dashboard
    df = df[df['Root_Cause'] != "Won't do"].copy()
    # Convert date column
    df['User_Reported_Date'] = pd.to_datetime(df['User_Reported_Date'], errors='coerce')
    return df


class FailureAnalysisTool:
    def __init__(self):
        self.df = None
        self.program_name = None
        
    def load_data(self, file, program_name):
        """Load CSV data and filter out 'Won't do' cases (cached parse)."""
        self.df = _load_dashboard_dataframe(file.getvalue())
        self.program_name = program_name
        return self.df
    
    def get_summary_stats(self):
        """Calculate summary statistics"""
        total_returns = len(self.df)
        root_cause_identified = len(self.df[self.df['Root_Cause'] == 'Root Cause Identified'])
        no_failure_found = len(self.df[self.df['Root_Cause'] == 'No Failure Found'])
        wont_do = len(self.df[self.df['Root_Cause'] == "Won't do"])
        
        return {
            'total_returns': total_returns,
            'root_cause_identified': root_cause_identified,
            'no_failure_found': no_failure_found,
            'wont_do': wont_do,
            'analysis_rate': f"{(root_cause_identified/total_returns*100):.1f}%" if total_returns > 0 else "0%"
        }
    
    def plot_return_reasons(self):
        """Plot return reason distribution"""
        reason_counts = self.df['Return_Reason_Code'].value_counts()
        if reason_counts.empty:
            return None
        
        fig = px.bar(
            x=reason_counts.index,
            y=reason_counts.values,
            labels={'x': 'Return Reason', 'y': 'Count'},
            title=f'{self.program_name} - Return Reasons Distribution',
            color=reason_counts.values,
            color_continuous_scale=CONTINUOUS_SCALE
        )
        fig.update_layout(showlegend=False, xaxis_tickangle=-45, coloraxis_showscale=False)
        return theme.style_fig(fig)
    
    def plot_root_cause_analysis(self):
        """Plot root cause breakdown"""
        root_cause_df = self.df[self.df['Root_Cause'].notna()]
        cause_counts = root_cause_df['Root_Cause'].value_counts()
        if cause_counts.empty:
            return None
        
        fig = px.pie(
            values=cause_counts.values,
            names=cause_counts.index,
            title=f'{self.program_name} - Root Cause Analysis Status',
            hole=0.55,
            color_discrete_sequence=CHART_SEQUENCE
        )
        fig.update_traces(marker=dict(line=dict(color=BG_ELEV, width=2)))
        return theme.style_fig(fig)
    
    def plot_timeline(self):
        """Plot returns over time"""
        timeline_df = self.df[self.df['User_Reported_Date'].notna()].copy()
        if timeline_df.empty:
            return None  # no report dates (e.g. Jupiter tracker) — skip the timeline
        timeline_df['Month'] = timeline_df['User_Reported_Date'].dt.to_period('M').astype(str)
        monthly_counts = timeline_df.groupby('Month').size().reset_index(name='Count')
        
        fig = px.line(
            monthly_counts,
            x='Month',
            y='Count',
            title=f'{self.program_name} - Returns Timeline',
            markers=True
        )
        fig.update_traces(line_color=PURPLE, line_width=3,
                          marker=dict(color=MAGENTA, size=8))
        return theme.style_fig(fig)
    
    def plot_sw_hw_breakdown(self):
        """Plot SW vs HW issues"""
        sw_yes = len(self.df[self.df['SW_Related_Issue'] == 'YES'])
        hw_yes = len(self.df[self.df['HW_Related_Issue'] == 'YES'])
        sw_no = len(self.df[self.df['SW_Related_Issue'] == 'NO'])
        hw_no = len(self.df[self.df['HW_Related_Issue'] == 'NO'])
        
        fig = go.Figure(data=[
            go.Bar(name='Software', x=['Related', 'Not Related'], y=[sw_yes, sw_no],
                   marker_color=PURPLE),
            go.Bar(name='Hardware', x=['Related', 'Not Related'], y=[hw_yes, hw_no],
                   marker_color=GREEN)
        ])
        fig.update_layout(
            title=f'{self.program_name} - SW vs HW Issue Distribution',
            barmode='group',
            xaxis_title='Issue Type',
            yaxis_title='Count'
        )
        return theme.style_fig(fig)
    
    def create_fault_tree(self):
        """Create fault wheel analysis using hierarchical structure"""
        # FWA Methodology: Top-down deductive analysis
        # Top Event -> Intermediate Events -> Basic Events (root causes)
        
        # Analyze root causes from data
        root_causes = self.df[self.df['Root_Cause_Reason'].notna()]['Root_Cause_Reason'].value_counts()
        
        if len(root_causes) == 0:
            # Return empty figure if no data
            fig = go.Figure()
            fig.add_annotation(text="No failure data available for fault wheel analysis",
                             xref="paper", yref="paper",
                             x=0.5, y=0.5, showarrow=False)
            return fig
        
        # Build hierarchical fault wheel structure
        # Level 0: Top Event (System Failure)
        # Level 1: Intermediate Events (Failure Categories)
        # Level 2: Basic Events (Root Causes)
        
        # Categorize root causes into intermediate events
        categories = {
            'Hardware Failures': [],
            'Software Failures': [],
            'Environmental Failures': [],
            'Power Failures': [],
            'Installation Failures': []
        }
        
        for cause, count in root_causes.items():
            cause_lower = str(cause).lower()
            
            # Categorize based on keywords
            if any(kw in cause_lower for kw in ['emmc', 'component', 'capacitor', 'solder', 'pcb', 'hardware', 'eipd', 'eos']):
                categories['Hardware Failures'].append((cause, count))
            elif any(kw in cause_lower for kw in ['firmware', 'software', 'cloud', 'registration', 'certificate']):
                categories['Software Failures'].append((cause, count))
            elif any(kw in cause_lower for kw in ['liquid', 'ingress', 'temperature', 'thermal', 'moisture', 'corrosion']):
                categories['Environmental Failures'].append((cause, count))
            elif any(kw in cause_lower for kw in ['poe', 'power', 'voltage', 'electrical']):
                categories['Power Failures'].append((cause, count))
            elif any(kw in cause_lower for kw in ['mount', 'bracket', 'installation', 'setup']):
                categories['Installation Failures'].append((cause, count))
            else:
                # Default to hardware
                categories['Hardware Failures'].append((cause, count))
        
        # Remove empty categories
        categories = {k: v for k, v in categories.items() if len(v) > 0}
        
        # Build tree structure for visualization
        labels = ['System Failure<br>(Top Event)']
        parents = ['']
        values = [root_causes.sum()]
        colors = [MAGENTA]  # Magenta for top event
        
        # Add intermediate events (categories)
        for category, causes in categories.items():
            category_total = sum(count for _, count in causes)
            labels.append(f'{category}<br>({category_total} cases)')
            parents.append('System Failure<br>(Top Event)')
            values.append(category_total)
            colors.append(PURPLE)  # Purple for intermediate
            
            # Add basic events (root causes)
            for cause, count in causes[:5]:  # Limit to top 5 per category
                labels.append(f'{cause}<br>({count})')
                parents.append(f'{category}<br>({category_total} cases)')
                values.append(count)
                colors.append(GREEN)  # Green for basic events
        
        # Create sunburst diagram for fault wheel
        # White labels read cleanly on the purple/green/magenta segments.
        text_colors = ['#ffffff'] * len(labels)

        fig = go.Figure(go.Sunburst(
            labels=labels,
            parents=parents,
            values=values,
            branchvalues="total",
            marker=dict(
                colors=colors,
                line=dict(color=BG_ELEV, width=2)
            ),
            hovertemplate='<b>%{label}</b><br>Cases: %{value}<br>%{percentParent}<extra></extra>',
            textfont=dict(size=12, color=text_colors),
            insidetextfont=dict(color=text_colors)
        ))
        
        fig.update_layout(
            title={
                'text': f'{self.program_name} - Fault Wheel Analysis (FWA)<br><sub>Top-Down Deductive Analysis: System Failure → Categories → Root Causes</sub>',
                'x': 0.5,
                'xanchor': 'center'
            },
            height=600,
            margin=dict(t=100, l=0, r=0, b=0),
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
            font=dict(family="DM Sans, sans-serif", color=TEXT)
        )
        
        return fig
    
    def plot_shipment_status(self):
        """Plot shipment status. Returns None when there's no data so the
        dashboard can skip it entirely."""
        status_counts = self.df['Shipment_Status'].value_counts()
        if status_counts.empty:
            return None
        
        fig = px.bar(
            x=status_counts.index,
            y=status_counts.values,
            title=f'{self.program_name} - Shipment Status',
            labels={'x': 'Status', 'y': 'Count'},
            color=status_counts.values,
            color_continuous_scale=CONTINUOUS_SCALE
        )
        fig.update_layout(coloraxis_showscale=False)
        return theme.style_fig(fig)
    
    def generate_report(self):
        """Generate Word document report"""
        doc = Document()
        
        # Title
        title = doc.add_heading(f'{self.program_name} Failure Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()
        
        # Summary Statistics
        doc.add_heading('Executive Summary', 1)
        stats = self.get_summary_stats()
        doc.add_paragraph(f"Total Returns: {stats['total_returns']}")
        doc.add_paragraph(f"Root Cause Identified: {stats['root_cause_identified']}")
        doc.add_paragraph(f"No Failure Found: {stats['no_failure_found']}")
        doc.add_paragraph(f"Won't Do: {stats['wont_do']}")
        doc.add_paragraph(f"Analysis Success Rate: {stats['analysis_rate']}")
        doc.add_paragraph()
        
        # Top Return Reasons
        doc.add_heading('Top Return Reasons', 1)
        top_reasons = self.df['Return_Reason_Code'].value_counts().head(10)
        for reason, count in top_reasons.items():
            doc.add_paragraph(f"{reason}: {count}", style='List Bullet')
        doc.add_paragraph()
        
        # Root Cause Analysis
        doc.add_heading('Root Cause Breakdown', 1)
        root_causes = self.df[self.df['Root_Cause_Reason'].notna()]['Root_Cause_Reason'].value_counts()
        for cause, count in root_causes.items():
            doc.add_paragraph(f"{cause}: {count}", style='List Bullet')
        doc.add_paragraph()
        
        # SW vs HW
        doc.add_heading('Software vs Hardware Issues', 1)
        sw_yes = len(self.df[self.df['SW_Related_Issue'] == 'YES'])
        hw_yes = len(self.df[self.df['HW_Related_Issue'] == 'YES'])
        doc.add_paragraph(f"Software Related: {sw_yes}")
        doc.add_paragraph(f"Hardware Related: {hw_yes}")
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

@st.cache_data(show_spinner=False)
def _generate_report_bytes(file_bytes, program_name):
    """Build the Word report once per (file, program). Returns raw bytes so the
    result is cacheable and reused across reruns for the download button."""
    tool = FailureAnalysisTool()
    tool.df = _load_dashboard_dataframe(file_bytes)
    tool.program_name = program_name
    return tool.generate_report().getvalue()


# Initialize app
def _render_program_selector():
    """Render the program selection landing page. Returns True if a program is selected."""
    theme.render_app_header()

    programs = get_program_list()

    if programs:
        st.markdown("### Select a Program")
        cols = st.columns(min(len(programs), 4))
        registry = load_registry()
        for i, name in enumerate(programs):
            info = registry["programs"].get(name, {})
            with cols[i % len(cols)]:
                st.markdown(
                    f'<div style="background:linear-gradient(160deg,{PANEL} 0%,{BG_ELEV} 100%);'
                    f'border:1px solid {BORDER};border-left:3px solid {PURPLE};border-radius:14px;'
                    f'padding:18px 16px;text-align:center;box-shadow:0 2px 16px #00000033;min-height:120px;">'
                    f'<div style="font-size:1.3em;font-weight:700;color:{TEXT};">{info.get("display_name", name)}</div>'
                    f'<div style="color:{PURPLE_BRIGHT};font-size:.85em;margin:6px 0;font-weight:600;">{info.get("product", "")}</div>'
                    f'<div style="color:{TEXT_MUTED};font-size:.8em;">{info.get("description", "")}</div></div>',
                    unsafe_allow_html=True,
                )
                if st.button(f"Open {name}", key=f"sel_{name}", use_container_width=True):
                    set_selected_program(name)
                    st.rerun()

    st.markdown("---")
    st.markdown("### Register New Program")
    with st.form("new_program_form"):
        c1, c2 = st.columns(2)
        with c1:
            new_name = st.text_input("Program Name", placeholder="e.g. Falcon")
            new_product = st.text_input("Product Name", placeholder="e.g. eero Indoor 7")
        with c2:
            new_display = st.text_input("Display Name (optional)", placeholder="Same as program name if blank")
            new_desc = st.text_input("Description", placeholder="e.g. WiFi 7 Indoor Mesh Router")
        submitted = st.form_submit_button("Register Program", use_container_width=True)
        if submitted and new_name.strip():
            register_program(
                new_name.strip(),
                display_name=new_display.strip() or None,
                product=new_product.strip() or None,
                description=new_desc.strip(),
            )
            set_selected_program(new_name.strip())
            st.rerun()
        elif submitted:
            st.warning("Please enter a program name.")


def main():
    # If no program selected yet, show the selector
    selected = get_selected_program()
    if not selected:
        _render_program_selector()
        return

    program_name = selected

    theme.render_app_header(f"Program: {program_name}")
    
    # Sidebar
    with st.sidebar:
        st.markdown(
            f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:6px;">'
            f'{theme.logo_svg(30)}<span style="font-weight:700;font-size:1.05em;color:{TEXT};">{APP_NAME}</span></div>',
            unsafe_allow_html=True,
        )
        st.header("⚙️ Configuration")

        # Show current program with option to switch
        st.markdown(f"**Program:** {program_name}")
        if st.button("↩ Switch Program", use_container_width=True):
            set_selected_program(None)
            # Clear program-specific session state
            for key in list(st.session_state.keys()):
                if key.startswith("debugger_") or key in ("triage_assistant", "vi_findings", "vi_notes"):
                    del st.session_state[key]
            st.rerun()
        
        # File upload
        uploaded_file = st.file_uploader(
            "Upload CSV File",
            type=['csv'],
            help="Upload your failure returns CSV file"
        )
        
        st.markdown("---")
        st.markdown("### 📑 View Selection")
        view_mode = st.radio(
            "Select View",
            options=["📊 Dashboard", "🔧 Triage Assistant", "🔬 Debug Workspace", "PCB Debugger",
                     "📈 Debug Analytics", "📋 Failure Analysis Table", "🔭 Data Explorer",
                     "📐 Schematic Viewer"],
            help="Debug Workspace = measurements and schematic side by side (recommended). "
                 "PCB Debugger = the classic guided flow. Schematic Viewer = browse/search "
                 "the sheets. Neither of the last three needs a CSV."
        )
        
        if view_mode == "📊 Dashboard":
            st.markdown("---")
            st.markdown("### 📊 Analysis Options")
            show_summary = st.checkbox("Summary Statistics", value=True)
            show_charts = st.checkbox("Visualizations", value=True)
            show_fault_tree = st.checkbox("Fault Wheel Analysis", value=True)
            show_data_table = st.checkbox("Data Table", value=False)
        else:
            show_summary = False
            show_charts = False
            show_fault_tree = False
            show_data_table = False
        
        # Exit button at the bottom
        st.markdown("---")
        if st.button("🚪 Exit Application", type="secondary", use_container_width=True):
            st.success("👋 Thank you for using the Failure Analysis Tool!")
            st.info("You can safely close this browser tab now.")
            st.balloons()
            st.stop()

        # Version footer
        st.markdown(
            f'<div style="text-align:center;color:{TEXT_FAINT};font-size:.75em;'
            f'font-family:\'DM Mono\',monospace;margin-top:14px;">{APP_NAME} · v{APP_VERSION}</div>',
            unsafe_allow_html=True,
        )
        
    # These views work on the program's own design assets, so they don't need
    # (and must not wait for) a field-return CSV upload.
    if view_mode == "📐 Schematic Viewer":
        render_schematic_viewer()
        return
    if view_mode == "🔬 Debug Workspace":
        render_debug_workspace()
        return

    # Main content
    if uploaded_file is not None:
        # The Data Explorer is schema-agnostic — always explore the raw upload.
        if view_mode == "🔭 Data Explorer":
            try:
                raw_df = pd.read_csv(io.BytesIO(uploaded_file.getvalue()))
            except Exception as e:
                st.error(f"Could not read the CSV: {e}")
                return
            render_data_explorer(raw_df, program_name)
            return

        # For the analysis views, auto-convert a Jupiter tracker export into the
        # failure-returns schema so the Dashboard / Fault Wheel work from the
        # ORIGINAL file (no manual conversion needed).
        file_bytes = uploaded_file.getvalue()
        try:
            _cols = set(pd.read_csv(io.BytesIO(file_bytes), nrows=1).columns)
            if convert_jupiter.looks_like_jupiter(_cols) and "Root_Cause" not in _cols:
                file_bytes = convert_jupiter.convert_bytes(file_bytes)
                st.success("Detected a Jupiter tracker CSV — auto-converted it to the "
                           "failure-analysis schema so the Dashboard and Fault Wheel work directly.")
        except Exception:
            pass  # fall through to the schema guard below

        # Schema guard: the Dashboard / Triage / Table / Analytics views require
        # the field-returns schema. If the CSV still doesn't have it, fall back
        # to the schema-agnostic Data Explorer instead of crashing.
        try:
            _peek = pd.read_csv(io.BytesIO(file_bytes), nrows=5)
        except Exception as e:
            st.error(f"Could not read the CSV: {e}")
            return
        _required = {"Root_Cause", "Return_Reason_Code", "User_Reported_Date"}
        _missing = _required - set(_peek.columns)
        if _missing:
            st.info(
                "This CSV doesn't match the failure-returns schema (missing columns: "
                f"**{', '.join(sorted(_missing))}**), so the *{view_mode}* view can't be built "
                "from it. Showing the **Data Explorer** instead — or pick the 🔭 Data Explorer "
                "view in the sidebar for arbitrary CSVs like this one."
            )
            raw_df = pd.read_csv(io.BytesIO(file_bytes))
            render_data_explorer(raw_df, program_name)
            return

        tool = FailureAnalysisTool()
        df = tool.load_data(io.BytesIO(file_bytes), program_name)
        
        # Route to appropriate view
        if view_mode == "🔧 Triage Assistant":
            render_triage_ui(df)
            return
        elif view_mode == "📋 Failure Analysis Table":
            render_failure_table(df, program_name)
            return
        elif view_mode == "PCB Debugger":
            render_debugger_ui()
            return
        elif view_mode == "📈 Debug Analytics":
            render_analytics_ui()
            return
        
        # Summary Statistics
        if show_summary:
            st.header("📈 Summary Statistics")
            stats = tool.get_summary_stats()
            
            col1, col2, col3, col4, col5 = st.columns(5)
            with col1:
                st.metric("Total Returns", stats['total_returns'])
            with col2:
                st.metric("Root Cause ID'd", stats['root_cause_identified'])
            with col3:
                st.metric("No Failure Found", stats['no_failure_found'])
            with col4:
                st.metric("Won't Do", stats['wont_do'])
            with col5:
                st.metric("Analysis Rate", stats['analysis_rate'])
            
            st.markdown("---")
        
        # Visualizations
        if show_charts:
            st.header("📊 Visualizations")

            def _safe_chart(fn, col=None):
                """Render a chart if it has data; never let one chart abort the page."""
                try:
                    fig = fn()
                except Exception as e:
                    (col or st).caption(f"(chart skipped: {e})")
                    return
                if fig is None:
                    return  # no data for this chart — skip it silently
                (col or st).plotly_chart(fig, use_container_width=True)

            # Row 1
            col1, col2 = st.columns(2)
            _safe_chart(tool.plot_return_reasons, col1)
            _safe_chart(tool.plot_root_cause_analysis, col2)

            # Row 2
            col1, col2 = st.columns(2)
            _safe_chart(tool.plot_timeline, col1)
            _safe_chart(tool.plot_sw_hw_breakdown, col2)

            # Row 3 (shipment status — skipped automatically when not tracked)
            _safe_chart(tool.plot_shipment_status)

            st.markdown("---")
        
        # Fault Wheel Analysis
        if show_fault_tree:
            st.header("Fault Wheel Analysis")
            try:
                st.plotly_chart(tool.create_fault_tree(), use_container_width=True)
            except Exception as e:
                st.warning(f"Fault wheel could not be rendered: {e}")
            st.markdown("---")
        
        # Data Table
        if show_data_table:
            st.header("📋 Data Table")
            st.dataframe(df, use_container_width=True)
            st.markdown("---")
        
        # Report Generation
        st.header("📄 Report Generation")
        # Build the report once per (file, program); cached across reruns so the
        # Word document isn't regenerated on every interaction.
        report_bytes = _generate_report_bytes(file_bytes, program_name)
        col1, col2, col3 = st.columns([1, 1, 2])
        with col1:
            if st.button("🔄 Generate Report", type="primary"):
                st.success("Report generated successfully!")

        with col2:
            st.download_button(
                label="⬇️ Download Report",
                data=report_bytes,
                file_name=f"{program_name}_Failure_Analysis_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        # Export filtered data
        with col3:
            csv_buffer = io.StringIO()
            df.to_csv(csv_buffer, index=False)
            st.download_button(
                label="⬇️ Export Data (CSV)",
                data=csv_buffer.getvalue(),
                file_name=f"{program_name}_data_export_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )
    
    else:
        # PCB Debugger works without CSV data
        if view_mode == "PCB Debugger":
            render_debugger_ui()
            return
        if view_mode == "📈 Debug Analytics":
            render_analytics_ui()
            return

        # Welcome screen
        st.info("Please upload a CSV file from the sidebar to begin analysis")
        st.markdown(f"""
        ### Welcome to {APP_NAME}
        
        This tool helps you analyze field return data and identify patterns in failures.
        
        **Features:**
        - 📊 Interactive visualizations and charts
        - Fault wheel analysis
        - 📈 Trend analysis over time
        - 🔍 SW vs HW issue breakdown
        - 📄 Automated report generation
        - 💾 Data export capabilities
        
        **Getting Started:**
        1. Enter your program name in the sidebar
        2. Upload your CSV file
        3. Explore the visualizations and insights
        4. Generate and download reports
        """)


def render_failure_table(df_original, program_name):
    """Render comprehensive failure analysis table"""
    st.header(f"📋 {program_name} - Comprehensive Failure Analysis Table")
    
    # Load original data (before filtering)
    df = df_original.copy()
    
    # Summary metrics
    col1, col2, col3, col4 = st.columns(4)
    
    total_returns = len(df)
    real_failures = len(df[df['Root_Cause'] == 'Root Cause Identified'])
    ntf_cases = len(df[df['Root_Cause'] == 'No Failure Found'])
    wont_do = len(df[df['Root_Cause'] == "Won't do"])
    
    col1.metric("Total Returns", total_returns)
    col2.metric("✅ Real Failures", real_failures, f"{real_failures/total_returns*100:.1f}%")
    col3.metric("❌ NTF", ntf_cases, f"{ntf_cases/total_returns*100:.1f}%")
    col4.metric("⚠️ Won't Do", wont_do, f"{wont_do/total_returns*100:.1f}%")
    
    st.markdown("---")
    
    # Real Failures
    with st.expander("✅ REAL FAILURES - Root Cause Identified", expanded=True):
        real_fail = df[df['Root_Cause'] == 'Root Cause Identified'].copy()
        st.write(f"**Total: {len(real_fail)} cases**")
        
        if len(real_fail) > 0:
            # Normalize root causes
            cause_summary = real_fail['Root_Cause_Reason'].value_counts()
            st.markdown("**Breakdown by Root Cause:**")
            for cause, count in cause_summary.items():
                st.write(f"- **{cause}**: {count} case(s)")
            
            st.markdown("**Detailed Cases:**")
            display_cols = ['ID', 'User_Reported_Date', 'Return_Reason_Code', 'Root_Cause_Reason', 
                           'Power_Adapter', 'SW_Related_Issue', 'HW_Related_Issue', 'Jira_Ticket']
            st.dataframe(real_fail[display_cols], use_container_width=True, hide_index=True)
    
    # No Trouble Found
    with st.expander("❌ NO TROUBLE FOUND (NTF)", expanded=True):
        ntf = df[df['Root_Cause'] == 'No Failure Found'].copy()
        st.write(f"**Total: {len(ntf)} cases**")
        
        if len(ntf) > 0:
            ntf_reasons = ntf['Return_Reason_Code'].value_counts()
            st.markdown("**Breakdown by Return Reason:**")
            for reason, count in ntf_reasons.items():
                st.write(f"- **{reason}**: {count} case(s)")
            
            st.markdown("**Detailed Cases:**")
            display_cols = ['ID', 'User_Reported_Date', 'Return_Reason_Code', 'Root_Cause_Reason', 'Comments']
            st.dataframe(ntf[display_cols], use_container_width=True, hide_index=True)
    
    # Won't Do
    with st.expander("⚠️ WON'T DO - Not Analyzed (Units Not Returned)", expanded=False):
        wont = df[df['Root_Cause'] == "Won't do"].copy()
        st.warning(f"**{len(wont)} cases** marked as 'Won't do' - units were never returned for analysis")
        
        if len(wont) > 0:
            wont_reasons = wont['Return_Reason_Code'].value_counts()
            st.markdown("**Breakdown by Return Reason:**")
            for reason, count in wont_reasons.items():
                st.write(f"- **{reason}**: {count} case(s)")
            
            st.markdown("**Detailed Cases:**")
            display_cols = ['ID', 'User_Reported_Date', 'Return_Reason_Code', 'Unit_SN', 'Comments']
            st.dataframe(wont[display_cols], use_container_width=True, hide_index=True)
    
    st.markdown("---")
    
    # DAA Analysis
    st.subheader("🔍 DAA (Dead After Arrival) Analysis")
    daa_cases = df[df['Return_Reason_Code'] == 'DAA'].copy()
    
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total DAA", len(daa_cases))
    col2.metric("Real Failures", len(daa_cases[daa_cases['Root_Cause'] == 'Root Cause Identified']))
    col3.metric("NTF", len(daa_cases[daa_cases['Root_Cause'] == 'No Failure Found']))
    col4.metric("Won't Do", len(daa_cases[daa_cases['Root_Cause'] == "Won't do"]))
    
    daa_real = daa_cases[daa_cases['Root_Cause'] == 'Root Cause Identified']
    if len(daa_real) > 0:
        st.markdown("**DAA Root Causes:**")
        for idx, row in daa_real.iterrows():
            st.write(f"- **ID {row['ID']}**: {row['Root_Cause_Reason']}")
    
    st.markdown("---")
    
    # Power/PSU Analysis
    st.subheader("⚡ Power Adapter / PSU Analysis")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Power Adapter Distribution:**")
        power_dist = df['Power_Adapter'].value_counts()
        for adapter, count in power_dist.head(10).items():
            if pd.notna(adapter):
                st.write(f"- {adapter}: {count}")
    
    with col2:
        st.markdown("**PSU-Related Failures:**")
        goldfinch_cases = df[df['Power_Adapter'] == 'Goldfinch']
        goldfinch_failures = goldfinch_cases[goldfinch_cases['Root_Cause'] == 'Root Cause Identified']
        
        st.write(f"- Goldfinch (30W PSU) total: {len(goldfinch_cases)}")
        st.write(f"- Goldfinch failures: {len(goldfinch_failures)}")
        
        exothermic = df[df['Root_Cause_Reason'].str.contains('exothermic', case=False, na=False)]
        if len(exothermic) > 0:
            st.write(f"- **Exothermic events: {len(exothermic)}**")
    
    st.markdown("---")
    
    # Liquid Ingress
    st.subheader("💧 Liquid Ingress Analysis")
    liquid_cases = df[df['Root_Cause_Reason'].str.contains('liquid|ingress', case=False, na=False)]
    
    st.write(f"**Total: {len(liquid_cases)} cases**")
    st.info("Note: Case variations (liquid ingress / Liquid Ingress) are the same failure mode")
    
    for idx, row in liquid_cases.iterrows():
        st.write(f"- **ID {row['ID']}**: {row['Root_Cause_Reason']}")


if __name__ == "__main__":
    main()

